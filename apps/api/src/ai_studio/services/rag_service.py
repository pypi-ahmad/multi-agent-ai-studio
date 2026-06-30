from __future__ import annotations

import asyncio
import math
import subprocess
import tempfile
from collections import deque
from dataclasses import dataclass
from pathlib import Path
from typing import Any, cast
from urllib.parse import urljoin, urlparse
from uuid import uuid4

import httpx
import pandas as pd
import trafilatura
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from pypdf import PdfReader
from qdrant_client import QdrantClient
from qdrant_client.http import models as qdrant_models
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from ai_studio.core.config import get_settings
from ai_studio.models.entities import Document, EmbeddingMetadata
from ai_studio.schemas.rag import ConnectorIngestRequest, DocumentCreate, RetrievalHit, RetrievalRequest, RetrievalResponse
from ai_studio.services.model_router import ModelRouter
from ai_studio.services.ollama_client import OllamaClient

_COLLECTION_NAME = "studio_chunks"
_TEXT_EXTENSIONS = {
    ".txt",
    ".md",
    ".rst",
    ".py",
    ".js",
    ".ts",
    ".tsx",
    ".json",
    ".yaml",
    ".yml",
    ".toml",
    ".ini",
    ".csv",
    ".sql",
    ".ipynb",
}


@dataclass(slots=True)
class RetrievalCandidate:
    chunk_id: str
    document_id: str
    text: str
    metadata: dict[str, object]
    semantic_score: float = 0.0
    keyword_score: float = 0.0
    rerank_score: float = 0.0


class RagService:
    """Production-oriented local RAG service with Qdrant + Ollama embeddings."""

    def __init__(self, ollama_client: OllamaClient, model_router: ModelRouter) -> None:
        self._ollama_client = ollama_client
        self._model_router = model_router
        settings = get_settings()
        self._qdrant = QdrantClient(url=settings.qdrant_url, timeout=settings.qdrant_timeout_seconds)
        self._http = httpx.AsyncClient(timeout=20)

    async def close(self) -> None:
        await self._http.aclose()

    def _ensure_collection(self, vector_size: int) -> None:
        collections = self._qdrant.get_collections().collections
        if any(collection.name == _COLLECTION_NAME for collection in collections):
            return
        self._qdrant.create_collection(
            collection_name=_COLLECTION_NAME,
            vectors_config=qdrant_models.VectorParams(size=vector_size, distance=qdrant_models.Distance.COSINE),
        )

    @staticmethod
    def _chunk_text(text: str, chunk_size: int = 1200, overlap: int = 180) -> list[str]:
        normalized = "\n".join(line.strip() for line in text.splitlines() if line.strip())
        chunks: list[str] = []
        index = 0
        while index < len(normalized):
            chunks.append(normalized[index : index + chunk_size])
            index += chunk_size - overlap
        return [chunk.strip() for chunk in chunks if chunk.strip()]

    @staticmethod
    def _extract_text(path: Path) -> str:
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            reader = PdfReader(str(path))
            return "\n".join((page.extract_text() or "") for page in reader.pages)
        if suffix in {".html", ".htm"}:
            soup = BeautifulSoup(path.read_text(encoding="utf-8", errors="ignore"), "lxml")
            return soup.get_text("\n", strip=True)
        if suffix == ".docx":
            doc = DocxDocument(str(path))
            return "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())
        if suffix == ".csv":
            frame = pd.read_csv(path)
            return frame.to_csv(index=False)
        if suffix == ".xlsx":
            workbook = pd.ExcelFile(path)
            parts: list[str] = []
            for sheet_name in workbook.sheet_names:
                sheet = workbook.parse(sheet_name)
                parts.append(f"# Sheet: {sheet_name}\n" + sheet.to_csv(index=False))
            return "\n\n".join(parts)
        return path.read_text(encoding="utf-8", errors="ignore")

    @staticmethod
    def _normalize_scores(values: list[float]) -> list[float]:
        if not values:
            return []
        low = min(values)
        high = max(values)
        if math.isclose(low, high):
            return [1.0 for _ in values]
        return [(value - low) / (high - low) for value in values]

    @staticmethod
    def _keyword_score(text: str, query_terms: list[str]) -> float:
        lowered = text.lower()
        if not lowered or not query_terms:
            return 0.0
        hits = sum(lowered.count(term) for term in query_terms)
        return hits / max(len(lowered.split()), 1)

    @staticmethod
    def _highlights(text: str, query_terms: list[str], max_snippets: int = 3) -> list[str]:
        lowered = text.lower()
        snippets: list[str] = []
        for term in query_terms:
            pos = lowered.find(term)
            if pos < 0:
                continue
            start = max(pos - 80, 0)
            end = min(pos + len(term) + 80, len(text))
            snippets.append(text[start:end].strip())
            if len(snippets) >= max_snippets:
                break
        return snippets

    @staticmethod
    def _cosine_similarity(a: list[float], b: list[float]) -> float:
        if not a or not b or len(a) != len(b):
            return 0.0
        dot = sum(x * y for x, y in zip(a, b, strict=False))
        mag_a = math.sqrt(sum(x * x for x in a))
        mag_b = math.sqrt(sum(y * y for y in b))
        if mag_a == 0.0 or mag_b == 0.0:
            return 0.0
        return dot / (mag_a * mag_b)

    @staticmethod
    def _as_int(value: object, default: int) -> int:
        if isinstance(value, bool):
            return int(value)
        if isinstance(value, int):
            return value
        if isinstance(value, float):
            return int(value)
        if isinstance(value, str):
            try:
                return int(value)
            except ValueError:
                return default
        return default

    @staticmethod
    def _as_bool(value: object, default: bool) -> bool:
        if isinstance(value, bool):
            return value
        if isinstance(value, str):
            lowered = value.strip().lower()
            if lowered in {"true", "1", "yes"}:
                return True
            if lowered in {"false", "0", "no"}:
                return False
            return default
        if isinstance(value, (int, float)):
            return bool(value)
        return default

    @staticmethod
    def _normalize_filter_value(value: object) -> bool | int | str:
        if isinstance(value, (bool, int, str)):
            return value
        if isinstance(value, float):
            return int(value) if value.is_integer() else str(value)
        return str(value)

    async def register_document(
        self,
        session: AsyncSession,
        owner_id: str,
        payload: DocumentCreate,
    ) -> Document:
        document = Document(
            owner_id=owner_id,
            name=payload.name,
            mime_type=payload.mime_type,
            source_uri=payload.source_uri,
            meta=payload.metadata,
            status="registered",
        )
        session.add(document)
        await session.commit()
        await session.refresh(document)
        return document

    async def _index_text(
        self,
        session: AsyncSession,
        owner_id: str,
        document: Document,
        text: str,
        metadata: dict[str, object],
    ) -> int:
        chunks = self._chunk_text(text)
        if not chunks:
            document.status = "empty"
            await session.commit()
            return 0

        embedding_model = await self._model_router.pick("embedding")
        first_vector = await self._ollama_client.embeddings(embedding_model, chunks[0])
        self._ensure_collection(len(first_vector))

        points: list[qdrant_models.PointStruct] = []
        metadata_rows: list[EmbeddingMetadata] = []

        for idx, chunk in enumerate(chunks):
            vector = first_vector if idx == 0 else await self._ollama_client.embeddings(embedding_model, chunk)
            vector_id = str(uuid4())
            points.append(
                qdrant_models.PointStruct(
                    id=vector_id,
                    vector=vector,
                    payload={
                        "document_id": document.id,
                        "chunk_index": idx,
                        "text": chunk,
                        "owner_id": owner_id,
                        "source_uri": document.source_uri,
                        "name": document.name,
                        **{f"meta_{key}": value for key, value in metadata.items()},
                    },
                )
            )
            metadata_rows.append(
                EmbeddingMetadata(
                    document_id=document.id,
                    chunk_index=idx,
                    vector_id=vector_id,
                    meta={"embedding_model": embedding_model, **metadata},
                )
            )

        self._qdrant.upsert(collection_name=_COLLECTION_NAME, points=points)
        session.add_all(metadata_rows)
        document.status = "ingested"
        await session.commit()
        return len(chunks)

    async def ingest_local_file(self, session: AsyncSession, owner_id: str, document_id: str, local_path: Path) -> int:
        stmt = select(Document).where(Document.owner_id == owner_id, Document.id == document_id)
        result = await session.execute(stmt)
        document = result.scalar_one_or_none()
        if not document:
            raise ValueError("Document not found")

        text = self._extract_text(local_path)
        return await self._index_text(
            session=session,
            owner_id=owner_id,
            document=document,
            text=text,
            metadata={"source_type": "file", "filename": local_path.name},
        )

    async def ingest_connector(
        self,
        session: AsyncSession,
        owner_id: str,
        payload: ConnectorIngestRequest,
    ) -> tuple[Document, int]:
        document = await self.register_document(
            session,
            owner_id,
            DocumentCreate(
                name=payload.name,
                mime_type="text/plain",
                source_uri=payload.source_uri,
                metadata={"connector": payload.connector, **payload.metadata},
            ),
        )

        if payload.connector == "github":
            text = await self._collect_github_text(payload.source_uri, payload.options)
        elif payload.connector == "web":
            text = await self._crawl_web_text(payload.source_uri, payload.options)
        else:
            raise ValueError(f"Unsupported connector: {payload.connector}")

        chunks = await self._index_text(
            session=session,
            owner_id=owner_id,
            document=document,
            text=text,
            metadata={"source_type": payload.connector, **payload.options, **payload.metadata},
        )
        return document, chunks

    async def _collect_github_text(self, source_uri: str, options: dict[str, object]) -> str:
        branch = str(options.get("branch", ""))
        max_files = self._as_int(options.get("max_files", 80), 80)
        max_file_bytes = self._as_int(options.get("max_file_bytes", 250_000), 250_000)

        with tempfile.TemporaryDirectory(prefix="rag-github-") as temp_dir:
            repo_dir = Path(temp_dir) / "repo"
            command = ["git", "clone", "--depth", "1", source_uri, str(repo_dir)]
            if branch:
                command = ["git", "clone", "--depth", "1", "--branch", branch, source_uri, str(repo_dir)]

            def _clone() -> None:
                subprocess.run(command, check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)

            await asyncio.to_thread(_clone)

            files = sorted(path for path in repo_dir.rglob("*") if path.is_file() and path.suffix.lower() in _TEXT_EXTENSIONS)
            blocks: list[str] = []
            for path in files[:max_files]:
                if path.stat().st_size > max_file_bytes:
                    continue
                relative = path.relative_to(repo_dir)
                content = path.read_text(encoding="utf-8", errors="ignore")
                if not content.strip():
                    continue
                blocks.append(f"# File: {relative}\n{content}")
            return "\n\n".join(blocks)

    async def _crawl_web_text(self, start_url: str, options: dict[str, object]) -> str:
        max_pages = self._as_int(options.get("max_pages", 8), 8)
        max_depth = self._as_int(options.get("max_depth", 1), 1)
        same_domain = self._as_bool(options.get("same_domain", True), True)

        parsed_root = urlparse(start_url)
        queue: deque[tuple[str, int]] = deque([(start_url, 0)])
        visited: set[str] = set()
        pages: list[str] = []

        while queue and len(pages) < max_pages:
            url, depth = queue.popleft()
            if url in visited:
                continue
            visited.add(url)

            try:
                response = await self._http.get(url)
                response.raise_for_status()
            except Exception:  # noqa: BLE001
                continue

            html = response.text
            extracted = trafilatura.extract(html) or BeautifulSoup(html, "lxml").get_text("\n", strip=True)
            if extracted.strip():
                pages.append(f"# URL: {url}\n{extracted}")

            if depth >= max_depth:
                continue

            soup = BeautifulSoup(html, "lxml")
            for link in soup.find_all("a", href=True):
                next_url = urljoin(url, str(link["href"]))
                parsed_next = urlparse(next_url)
                if parsed_next.scheme not in {"http", "https"}:
                    continue
                if same_domain and parsed_next.netloc != parsed_root.netloc:
                    continue
                if next_url not in visited:
                    queue.append((next_url, depth + 1))

        return "\n\n".join(pages)

    def _build_owner_filter(self, owner_id: str, filters: dict[str, object]) -> qdrant_models.Filter:
        conditions: list[qdrant_models.Condition] = [
            qdrant_models.FieldCondition(key="owner_id", match=qdrant_models.MatchValue(value=owner_id))
        ]
        for key, value in filters.items():
            conditions.append(
                qdrant_models.FieldCondition(
                    key=key, match=qdrant_models.MatchValue(value=self._normalize_filter_value(value))
                )
            )
        return qdrant_models.Filter(must=conditions)

    def _keyword_candidates(
        self,
        owner_id: str,
        query_terms: list[str],
        filters: dict[str, object],
        limit: int,
    ) -> list[RetrievalCandidate]:
        owner_filter = self._build_owner_filter(owner_id, filters)
        records, _ = self._qdrant.scroll(
            collection_name=_COLLECTION_NAME,
            scroll_filter=owner_filter,
            with_payload=True,
            limit=max(limit * 4, 50),
        )
        candidates: list[RetrievalCandidate] = []
        for record in records:
            payload = cast(dict[str, Any], record.payload or {})
            text = str(payload.get("text", ""))
            score = self._keyword_score(text, query_terms)
            if score <= 0:
                continue
            candidates.append(
                RetrievalCandidate(
                    chunk_id=str(record.id),
                    document_id=str(payload.get("document_id", "")),
                    text=text,
                    metadata={
                        "chunk_index": payload.get("chunk_index"),
                        "source_uri": payload.get("source_uri", ""),
                        "name": payload.get("name", ""),
                    },
                    keyword_score=score,
                )
            )
        candidates.sort(key=lambda item: item.keyword_score, reverse=True)
        return candidates[:limit]

    async def retrieve(self, owner_id: str, payload: RetrievalRequest) -> RetrievalResponse:
        query_terms = [term for term in payload.query.lower().split() if len(term) > 1]
        embedding_model = await self._model_router.pick("embedding")
        query_vector = await self._ollama_client.embeddings(embedding_model, payload.query)
        owner_filter = self._build_owner_filter(owner_id, payload.filters)

        candidates: dict[str, RetrievalCandidate] = {}

        if payload.mode in {"semantic", "hybrid"}:
            semantic_hits = self._qdrant.query_points(
                collection_name=_COLLECTION_NAME,
                query=query_vector,
                query_filter=owner_filter,
                limit=payload.candidate_pool,
                with_payload=True,
            ).points
            for hit in semantic_hits:
                chunk_id = str(hit.id)
                payload_data = hit.payload or {}
                candidate = candidates.get(chunk_id)
                if candidate is None:
                    candidate = RetrievalCandidate(
                        chunk_id=chunk_id,
                        document_id=str(payload_data.get("document_id", "")),
                        text=str(payload_data.get("text", "")),
                        metadata={
                            "chunk_index": payload_data.get("chunk_index"),
                            "source_uri": payload_data.get("source_uri", ""),
                            "name": payload_data.get("name", ""),
                        },
                    )
                    candidates[chunk_id] = candidate
                candidate.semantic_score = float(hit.score)

        if payload.mode in {"keyword", "hybrid"}:
            for candidate in self._keyword_candidates(owner_id, query_terms, payload.filters, payload.candidate_pool):
                existing = candidates.get(candidate.chunk_id)
                if existing is None:
                    candidates[candidate.chunk_id] = candidate
                else:
                    existing.keyword_score = max(existing.keyword_score, candidate.keyword_score)

        ranked = list(candidates.values())
        if not ranked:
            return RetrievalResponse(query=payload.query, mode=payload.mode, hits=[])

        semantic_scores = self._normalize_scores([item.semantic_score for item in ranked])
        keyword_scores = self._normalize_scores([item.keyword_score for item in ranked])

        for index, item in enumerate(ranked):
            semantic = semantic_scores[index]
            keyword = keyword_scores[index]
            if payload.mode == "semantic":
                item.rerank_score = semantic
            elif payload.mode == "keyword":
                item.rerank_score = keyword
            else:
                item.rerank_score = 0.7 * semantic + 0.3 * keyword

        ranked.sort(key=lambda item: item.rerank_score, reverse=True)

        if payload.rerank:
            rerank_pool = ranked[: min(len(ranked), max(payload.top_k * 3, 12))]
            for item in rerank_pool:
                candidate_vector = await self._ollama_client.embeddings(embedding_model, item.text[:3500])
                similarity = self._cosine_similarity(query_vector, candidate_vector)
                item.rerank_score = 0.55 * item.rerank_score + 0.45 * max(similarity, 0.0)
            ranked = rerank_pool + ranked[len(rerank_pool) :]
            ranked.sort(key=lambda item: item.rerank_score, reverse=True)

        results = [
            RetrievalHit(
                document_id=item.document_id,
                chunk_id=item.chunk_id,
                score=round(item.rerank_score, 6),
                text=item.text,
                highlights=self._highlights(item.text, query_terms),
                metadata={
                    **item.metadata,
                    "semantic_score": round(item.semantic_score, 6),
                    "keyword_score": round(item.keyword_score, 6),
                    "embedding_model": embedding_model,
                },
            )
            for item in ranked[: payload.top_k]
        ]
        return RetrievalResponse(query=payload.query, mode=payload.mode, hits=results)
